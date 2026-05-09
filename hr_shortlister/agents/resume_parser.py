from __future__ import annotations

import io
import re
from pathlib import Path

from hr_shortlister.core.llm_client import LLMQuotaError, LLMResponseError, generate_structured_json
from hr_shortlister.core.prompts import RESUME_SYSTEM_PROMPT, resume_user_prompt
from hr_shortlister.core.schemas import CandidateProfile, CandidateSource, Project
from hr_shortlister.core.utils import bounded_text_block, sanitize_llm_input


def parse_resume_file(filename: str, data: bytes) -> CandidateProfile:
    raw_text = extract_resume_text(filename, data)
    cleaned = sanitize_llm_input(raw_text, max_chars=50000)
    if not cleaned:
        raise ValueError(f"No readable text found in {filename}.")

    prompt_text = bounded_text_block("RESUME", cleaned)
    try:
        profile = generate_structured_json(
            CandidateProfile,
            RESUME_SYSTEM_PROMPT,
            resume_user_prompt(prompt_text, Path(filename).stem),
        )
    except (LLMQuotaError, LLMResponseError):
        profile = parse_resume_text_locally(cleaned, filename)

    profile.source = CandidateSource.RESUME
    if not profile.raw_text:
        profile.raw_text = raw_text
    if not profile.name or profile.name == "Unknown Candidate":
        profile.name = Path(filename).stem
    return profile


def extract_resume_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf_text(data)
    if suffix == ".docx":
        return _extract_docx_text(data)

    raise ValueError(f"Unsupported resume type for {filename}. Use PDF or DOCX.")


def parse_resume_text_locally(text: str, filename: str) -> CandidateProfile:
    return CandidateProfile(
        source=CandidateSource.RESUME,
        name=_extract_name(text, Path(filename).stem),
        email=_extract_email(text),
        phone=_extract_phone(text),
        location=None,
        education=[],
        experience_years=_extract_experience_years(text),
        skills=_extract_known_skills(text),
        certifications=_extract_certifications(text),
        projects=_extract_projects(text),
        raw_text=text,
    )


def _extract_pdf_text(data: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required to read PDF resumes.") from exc

    with fitz.open(stream=data, filetype="pdf") as document:
        return "\n".join(page.get_text("text") for page in document)


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read DOCX resumes.") from exc

    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))
    return "\n".join(paragraphs + table_text)


def _extract_name(text: str, fallback: str) -> str:
    for line in text.splitlines()[:8]:
        cleaned = line.strip()
        if 2 <= len(cleaned.split()) <= 4 and not any(char.isdigit() for char in cleaned):
            if "resume" not in cleaned.lower() and "curriculum" not in cleaned.lower():
                return cleaned
    return fallback


def _extract_email(text: str) -> str | None:
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    return match.group(0) if match else None


def _extract_phone(text: str) -> str | None:
    match = re.search(r"(?:\+?\d[\d\s().-]{8,}\d)", text)
    return match.group(0).strip() if match else None


def _extract_experience_years(text: str) -> float:
    match = re.search(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)", text, flags=re.IGNORECASE)
    return float(match.group(1)) if match else 0.0


def _extract_known_skills(text: str) -> list[str]:
    known_skills = [
        "Python",
        "Java",
        "JavaScript",
        "TypeScript",
        "React",
        "Node.js",
        "FastAPI",
        "Django",
        "Flask",
        "SQL",
        "PostgreSQL",
        "MongoDB",
        "AWS",
        "Azure",
        "GCP",
        "Docker",
        "Kubernetes",
        "LangChain",
        "LangGraph",
        "RAG",
        "LLM",
        "LLMs",
        "Machine Learning",
        "Deep Learning",
        "NLP",
        "Qdrant",
        "Pinecone",
        "Streamlit",
        "Git",
    ]
    lowered = text.lower()
    return [skill for skill in known_skills if skill.lower() in lowered]


def _extract_certifications(text: str) -> list[str]:
    certifications: list[str] = []
    for line in text.splitlines():
        cleaned = line.strip(" -•\t")
        if re.search(r"\b(certified|certification|certificate)\b", cleaned, re.IGNORECASE):
            certifications.append(cleaned[:120])
    return certifications[:8]


def _extract_projects(text: str) -> list[Project]:
    projects: list[Project] = []
    lines = [line.strip(" -•\t") for line in text.splitlines() if line.strip()]
    for index, line in enumerate(lines):
        if "project" in line.lower() and len(line) < 100:
            description = lines[index + 1] if index + 1 < len(lines) else None
            projects.append(Project(name=line[:80], description=description, tech_stack=[]))
        if len(projects) >= 5:
            break
    if not projects and len(text) > 40:
        projects.append(Project(name="Resume Portfolio", description=text[:500], tech_stack=_extract_known_skills(text)))
    return projects
